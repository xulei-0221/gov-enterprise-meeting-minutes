"""
生成符合 GB/T 9704-2012 的党政机关红头公文 Word 模板。
输出：红头公文模板.docx

核心参数（GB/T 9704-2012）：
  - 用纸：A4（210 × 297 mm）
  - 页边距：上 37mm，下 35mm，左 28mm，右 26mm
  - 版心：156 × 225 mm（每页 22 行，每行 28 字）
  - 字体：份号/密级/紧急 三号黑体；发文机关标志 二号小标宋（红）；
          发文字号 三号仿宋_GB2312；标题 二号小标宋；正文 三号仿宋_GB2312
  - 行距：固定值 28 磅
  - 首行缩进：2 字符
  - 反线：红色武文线，宽度等同版心
"""

from docx import Document
from docx.shared import Pt, Mm, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement


RED = RGBColor(0xC0, 0x00, 0x00)


def set_run_font(run, font_name, size_pt, bold=False, color=None):
    """统一设置中英文字体、字号、颜色。"""
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rFonts.set(qn('w:eastAsia'), font_name)
    rFonts.set(qn('w:cs'), font_name)


def set_para_format(p, *, line_spacing_pt=28, first_line_chars=0,
                    align=None, space_before=0, space_after=0):
    """设置段落格式：固定行距 28 磅、首行缩进、对齐、段前段后。"""
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = Pt(line_spacing_pt)
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    if align is not None:
        p.alignment = align
    if first_line_chars:
        pPr = p._p.get_or_add_pPr()
        ind = pPr.find(qn('w:ind'))
        if ind is None:
            ind = OxmlElement('w:ind')
            pPr.append(ind)
        ind.set(qn('w:firstLineChars'), str(first_line_chars * 100))
        ind.set(qn('w:firstLine'), str(first_line_chars * 200))


def add_horizontal_rule(paragraph, color_hex='C00000', size=12):
    """在段落底部添加红色反线（武文线）。"""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), str(size))
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)


def set_doc_default_font(doc, font_name='仿宋_GB2312', size_pt=16):
    """设置全文默认字体（正文：三号 = 16 磅 仿宋_GB2312）。"""
    styles = doc.styles
    normal = styles['Normal']
    normal.font.name = font_name
    normal.font.size = Pt(size_pt)
    rPr = normal.element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rFonts.set(qn('w:eastAsia'), font_name)
    rFonts.set(qn('w:cs'), font_name)


def set_page(section):
    """A4 页面 + GB/T 9704 页边距。"""
    section.page_height = Mm(297)
    section.page_width = Mm(210)
    section.top_margin = Mm(37)
    section.bottom_margin = Mm(35)
    section.left_margin = Mm(28)
    section.right_margin = Mm(26)
    section.header_distance = Mm(15)
    section.footer_distance = Mm(20)


def add_page_number_footer(section):
    """页码居中，三号宋体，按 GB/T 9704-2012 在版心下边缘之下空一行。"""
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_format(p, line_spacing_pt=28)
    run = p.add_run('— ')
    set_run_font(run, '宋体', 14)
    fld = OxmlElement('w:fldSimple')
    fld.set(qn('w:instr'), 'PAGE')
    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), '28')
    rPr.append(sz)
    r.append(rPr)
    t = OxmlElement('w:t')
    t.text = '1'
    r.append(t)
    fld.append(r)
    p._p.append(fld)
    run2 = p.add_run(' —')
    set_run_font(run2, '宋体', 14)


def build_template(output_path):
    doc = Document()

    set_doc_default_font(doc, '仿宋_GB2312', 16)
    section = doc.sections[0]
    set_page(section)
    add_page_number_footer(section)

    # ─────────────── 版头区 ───────────────
    # 1. 份号（三号黑体，顶格左上角）
    p = doc.add_paragraph()
    set_para_format(p, line_spacing_pt=28, align=WD_ALIGN_PARAGRAPH.LEFT)
    run = p.add_run('〔份号请填写6位阿拉伯数字，非涉密删除本行〕')
    set_run_font(run, '黑体', 16)

    # 2. 密级与保密期限（三号黑体）
    p = doc.add_paragraph()
    set_para_format(p, line_spacing_pt=28, align=WD_ALIGN_PARAGRAPH.LEFT)
    run = p.add_run('〔密级★保密期限，如：机密★1年；非涉密删除本行〕')
    set_run_font(run, '黑体', 16)

    # 3. 紧急程度(三号黑体)
    p = doc.add_paragraph()
    set_para_format(p, line_spacing_pt=28, align=WD_ALIGN_PARAGRAPH.LEFT)
    run = p.add_run('〔紧急程度，如：特急/加急；非紧急删除本行〕')
    set_run_font(run, '黑体', 16)

    # 空两行
    for _ in range(2):
        p = doc.add_paragraph()
        set_para_format(p, line_spacing_pt=28)

    # 4. 发文机关标志（二号小标宋，红色，居中）
    p = doc.add_paragraph()
    set_para_format(p, line_spacing_pt=44, align=WD_ALIGN_PARAGRAPH.CENTER,
                    space_before=12, space_after=12)
    run = p.add_run('XX机关文件')
    set_run_font(run, '方正小标宋简体', 22, color=RED)

    # 空一行
    p = doc.add_paragraph()
    set_para_format(p, line_spacing_pt=28)

    # 5. 发文字号 + 签发人（三号仿宋_GB2312）
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    table.columns[0].width = Mm(78)
    table.columns[1].width = Mm(78)
    cell_l = table.cell(0, 0)
    cell_r = table.cell(0, 1)

    p = cell_l.paragraphs[0]
    set_para_format(p, line_spacing_pt=28, align=WD_ALIGN_PARAGRAPH.LEFT)
    run = p.add_run('〔发文字号，如：市府发〔2026〕1号〕')
    set_run_font(run, '仿宋_GB2312', 16)

    p = cell_r.paragraphs[0]
    set_para_format(p, line_spacing_pt=28, align=WD_ALIGN_PARAGRAPH.RIGHT)
    run = p.add_run('签发人：〔姓名，仅上行文标注〕')
    set_run_font(run, '仿宋_GB2312', 16)

    for cell in (cell_l, cell_r):
        tcPr = cell._tc.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')
        for edge in ('top', 'left', 'bottom', 'right'):
            b = OxmlElement(f'w:{edge}')
            b.set(qn('w:val'), 'nil')
            tcBorders.append(b)
        tcPr.append(tcBorders)

    # 6. 红色反线（武文线）—— 用一段空段落底部边框实现
    p = doc.add_paragraph()
    set_para_format(p, line_spacing_pt=12)
    add_horizontal_rule(p, color_hex='C00000', size=12)

    # ─────────────── 主体区 ───────────────
    # 7. 标题（二号小标宋，居中，回行对称）
    p = doc.add_paragraph()
    set_para_format(p, line_spacing_pt=44, align=WD_ALIGN_PARAGRAPH.CENTER,
                    space_before=18, space_after=18)
    run = p.add_run('〔会议名称〕纪要')
    set_run_font(run, '方正小标宋简体', 22)

    # 8. 主送机关（三号仿宋，顶格）— 纪要可省略，但保留占位
    p = doc.add_paragraph()
    set_para_format(p, line_spacing_pt=28, align=WD_ALIGN_PARAGRAPH.LEFT)
    run = p.add_run('〔主送机关，纪要类公文可省略本行〕')
    set_run_font(run, '仿宋_GB2312', 16)

    # 9. 会议基本信息（三号仿宋，左空两字非首段缩进）
    info_lines = [
        '会议时间：〔YYYY年M月D日 上午/下午 HH:MM-HH:MM〕',
        '会议地点：〔具体地点全称〕',
        '会议形式：〔现场会议 / 视频会议 / 现场+视频〕',
        '主持人：〔姓名（职务全称）〕',
        '出席人员：〔按职务从高到低排列；同单位集中、不同单位换行〕',
        '列席人员：〔上级单位代表、特邀嘉宾等；无则删除本行〕',
        '请假人员：〔如有；无则删除本行〕',
        '记录人：〔姓名（所属部门及职务）〕',
    ]
    for line in info_lines:
        p = doc.add_paragraph()
        set_para_format(p, line_spacing_pt=28, first_line_chars=2)
        run = p.add_run(line)
        set_run_font(run, '仿宋_GB2312', 16)

    # 空一行
    p = doc.add_paragraph()
    set_para_format(p, line_spacing_pt=28)

    # 10. 一级标题"一、会议议题"
    def add_h1(text):
        p = doc.add_paragraph()
        set_para_format(p, line_spacing_pt=28, first_line_chars=2,
                        align=WD_ALIGN_PARAGRAPH.LEFT)
        run = p.add_run(text)
        set_run_font(run, '黑体', 16)
        return p

    def add_h2(text):
        p = doc.add_paragraph()
        set_para_format(p, line_spacing_pt=28, first_line_chars=2)
        run = p.add_run(text)
        set_run_font(run, '楷体_GB2312', 16)
        return p

    def add_body(text):
        p = doc.add_paragraph()
        set_para_format(p, line_spacing_pt=28, first_line_chars=2)
        run = p.add_run(text)
        set_run_font(run, '仿宋_GB2312', 16)
        return p

    add_h1('一、会议议题')
    add_body('（一）〔议题1〕；')
    add_body('（二）〔议题2〕；')
    add_body('（三）〔议题3〕。')

    add_h1('二、会议内容')
    add_h2('（一）〔对应议题1的标题〕')
    add_body('〔发言整理：忠实于发言人观点，删除口头禅与重复，'
             '不夹带评价。可分点 1. 2. 3. 列出要点〕')
    add_h2('（二）〔对应议题2的标题〕')
    add_body('1. 工作成效：〔用数据和事实陈述〕')
    add_body('2. 存在问题：〔聚焦问题本身，不带情绪化评价〕')
    add_h2('（三）〔对应议题3的标题〕')
    add_body('1. 〔重点任务1〕')
    add_body('  （1）〔具体措施〕；')
    add_body('  （2）〔具体措施〕。')

    add_h1('三、会议要求')
    add_body('（一）〔要求1，建议采用"提高政治站位""强化责任落实"等动宾结构〕：〔展开说明〕。')
    add_body('（二）〔要求2〕：〔展开说明〕。')
    add_body('（三）〔要求3〕：〔展开说明〕。')

    add_h1('四、会议决议')
    add_body('会议经讨论，作出如下决议：')
    add_body('（一）〔决议事项1，含责任主体、完成时限、验收标准〕。')
    add_body('（二）〔决议事项2〕。')

    # 空一行
    p = doc.add_paragraph()
    set_para_format(p, line_spacing_pt=28)

    # 散会时间
    p = doc.add_paragraph()
    set_para_format(p, line_spacing_pt=28, first_line_chars=2)
    run = p.add_run('散会时间：〔HH:MM〕')
    set_run_font(run, '仿宋_GB2312', 16)

    # 空一行
    p = doc.add_paragraph()
    set_para_format(p, line_spacing_pt=28)

    # 11. 落款：发文机关 + 成文日期（右空四字）
    p = doc.add_paragraph()
    set_para_format(p, line_spacing_pt=28, align=WD_ALIGN_PARAGRAPH.RIGHT)
    run = p.add_run('〔发文机关全称〕    ')
    set_run_font(run, '仿宋_GB2312', 16)

    p = doc.add_paragraph()
    set_para_format(p, line_spacing_pt=28, align=WD_ALIGN_PARAGRAPH.RIGHT)
    run = p.add_run('〔YYYY年M月D日〕    ')
    set_run_font(run, '仿宋_GB2312', 16)

    # ─────────────── 版记区 ───────────────
    # 空一行
    p = doc.add_paragraph()
    set_para_format(p, line_spacing_pt=28)

    p = doc.add_paragraph()
    set_para_format(p, line_spacing_pt=28, align=WD_ALIGN_PARAGRAPH.LEFT)
    run = p.add_run('抄送：〔抄送机关全称或规范化简称；无则删除本行〕')
    set_run_font(run, '仿宋_GB2312', 14)

    p = doc.add_paragraph()
    set_para_format(p, line_spacing_pt=28, align=WD_ALIGN_PARAGRAPH.LEFT)
    add_horizontal_rule(p, color_hex='000000', size=6)

    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    table.columns[0].width = Mm(100)
    table.columns[1].width = Mm(56)
    cell_l = table.cell(0, 0)
    cell_r = table.cell(0, 1)

    p = cell_l.paragraphs[0]
    set_para_format(p, line_spacing_pt=28)
    run = p.add_run('〔印发机关全称，一般为发文机关办公厅（室）〕')
    set_run_font(run, '仿宋_GB2312', 14)

    p = cell_r.paragraphs[0]
    set_para_format(p, line_spacing_pt=28, align=WD_ALIGN_PARAGRAPH.RIGHT)
    run = p.add_run('〔YYYY年M月D日〕印发')
    set_run_font(run, '仿宋_GB2312', 14)

    for cell in (cell_l, cell_r):
        tcPr = cell._tc.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')
        for edge in ('top', 'left', 'bottom', 'right'):
            b = OxmlElement(f'w:{edge}')
            b.set(qn('w:val'), 'nil')
            tcBorders.append(b)
        tcPr.append(tcBorders)

    p = doc.add_paragraph()
    set_para_format(p, line_spacing_pt=28)
    add_horizontal_rule(p, color_hex='000000', size=6)

    doc.save(output_path)
    print(f'OK -> {output_path}')


if __name__ == '__main__':
    build_template('/Users/evan/Documents/trae_projects/会议纪要/红头公文模板.docx')
